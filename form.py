import streamlit as st
from docx import Document
from num2words import num2words
from mail import MonEmail
from datetime import datetime
import streamlit.components.v1 as components
import os

def main():
    # Paramètres SMTP
    serveur_smtp = 'smtp.office365.com'
    port_smtp = 587
    utilisateur = 'aiformulas@outlook.fr'
    password = os.getenv('PASSWORD')
    # Création de l'objet MonEmail avec les paramètres SMTP
    email_sender = MonEmail(serveur_smtp, port_smtp, utilisateur, password)

    # Paramètres de l'e-mail
    # destinataire = 'juridique.a2f@gmail.com'
    destinataire = 'dirilnoel1@gmail.com'
    sujet = 'Formulaire création société : '
    message = """
Bonjour,

Vous trouverez ci-joints, les documents necessaires à la création de la sosciété.

Cordialement,
AIFormulas
    """

    components.html("""
        <script>
        const inputs = window.parent.document.querySelectorAll('input');
        inputs.forEach(input => {
            input.addEventListener('keydown', function(event) {
                if (event.key === 'Enter') {
                    event.preventDefault();
                }
            });
        });
        </script>
    """,
    height=0)

    st.markdown("<h1 style='text-align: center;'>A2F AUDIT EXPERTISE CONSEIL</h1>", unsafe_allow_html=True)
    # Affichage du texte au-dessus des boutons radio
    st.markdown("<h3 style='text-align: center;'>Choisissez votre type de société :</h3>", unsafe_allow_html=True)

    # Liste des options pour les boutons radio
    options = ["SASU","SAS (2 actionnaires)", "SAS (3 actionnaires)"]

    # Affichage des boutons radio horizontalement
    selected_option = st.radio("", options, horizontal=True)

     # Récupérer la date actuelle
    aujourd_hui = datetime.now()
    date_aujourd_hui = aujourd_hui.strftime("%d/%m/%Y")

    # Affichage du formulaire en fonction de l'option sélectionnée
    if selected_option == "SAS (2 actionnaires)":
        st.header("Société par actions simplifiée (2 actionnaires)")
        with st.form(key='sas2', clear_on_submit=True):
            col1, col2 = st.columns(2)
            with col1:
                denominationsas2 = st.text_input("Nom de la société : ", key="denominationsas2")
                capitalsas2 = st.text_input("Capital social : ", key="capitalsas2",on_change=None)
                annee_bilansas2 = st.text_input("Année de clôture du premier bilan : ", key="annee_bilansas2")

                st.markdown("**Président associé 1**")
                sexe_associe_1sas2 = st.radio("Sexe", key="sexe_associe_1sas2", options=["Monsieur", "Madame"])
                prenom_associe_1sas2 = st.text_input("Prénom : ", key="prenom_associe_1sas2")
                nom_associe_1sas2 = st.text_input("Nom : ", key="nom_associe_1sas2")
                date_associe_1sas2 = st.text_input("Date de naissance(JJ/MM/AAAA) : ", key="date_associe_1sas2")
                lieu_naissance_associe_1sas2 = st.text_input("Lieu de naissance : ", key="lieu_naissance_associe_1sas2")
                code_postal_associe_1sas2 = st.text_input("Departement de naissance : ", key="code_postal_associe_1sas2")
                nationalite_associe_1sas2 = st.text_input("Nationalité : ", key="nationalite_associe_1sas2")
                adresse_associe_1sas2 = st.text_input("Adresse de domicile (adresse complète) :", key="adresse_associe_1sas2")
                apport_associe_1sas2 = st.text_input("Apport : ", key="apport_associe_1sas2")
                pere_associe_1sas2 = st.text_input("Nom et prénom du père : ", key="pere_associe_1sas2")
                mere_associe_1sas2 = st.text_input("Nom et prénom de la mère : ", key="mere_associe_1sas2")

                st.markdown('#')
                st.markdown('Vos informations')


                nomsas2 = st.text_input("Nom et prénom :", key="nom")
                emailsas2 = st.text_input("Email :", key="email")
                phonesas2 = st.text_input("Telephone :", key="phone")





            with col2:
                activitesas2 = st.text_input("Activité de la société : ", key="activitesas2")
                siege_socialsas2 = st.text_input("Siège social (adresse complète) : ", key="siege_socialsas2")
                ville_siege_socialsas2 = st.text_input("Ville du siège social : ", key="ville_siege_socialsas2")

                st.markdown("**Associé 2**")
                sexe_associe_2sas2 = st.radio("Sexe ", key="sexe_associe_2sas2", options=["Monsieur", "Madame"])
                prenom_associe_2sas2 = st.text_input("Prénom : ", key="prenom_associe_2sas2")
                nom_associe_2sas2 = st.text_input("Nom : ", key="nom_associe_2sas2")
                date_associe_2sas2 = st.text_input("Date de naissance (JJ/MM/AAAA) : ", key="date_associe_2sas2")
                lieu_naissance_associe_2sas2 = st.text_input("Lieu de naissance : ", key="lieu_naissance_associe_2sas2")
                code_postal_associe_2sas2 = st.text_input("Departement de naissance :", key="code_postal_associe_2sas2")
                nationalite_associe_2sas2 = st.text_input("Nationalité : ", key="nationalite_associe_2sas2")
                adresse_associe_2sas2 = st.text_input("Adresse de domicile (adresse complète) :", key="adresse_associe_2sas2")
                apport_associe_2sas2 = st.text_input("Apport : ", key="apport_associe_2sas2")


            submittedsas2 = st.form_submit_button("Valider")

            if submittedsas2:
                if all([denominationsas2, capitalsas2, activitesas2, siege_socialsas2, prenom_associe_1sas2, nom_associe_1sas2, date_associe_1sas2,
                        lieu_naissance_associe_1sas2, code_postal_associe_1sas2, nationalite_associe_1sas2, adresse_associe_1sas2,
                        annee_bilansas2, prenom_associe_2sas2, nom_associe_2sas2, date_associe_2sas2, lieu_naissance_associe_2sas2,
                        code_postal_associe_2sas2, nationalite_associe_2sas2, adresse_associe_2sas2, apport_associe_1sas2,apport_associe_2sas2,ville_siege_socialsas2,pere_associe_1sas2]):
                    
                    remplacements_sas2 = {
                        '<CAPITAL>': capitalsas2,
                        '<SIEGESOCIAL>': siege_socialsas2,
                        '<DENOMINATION>': denominationsas2.upper(),
                        '<SEXEASSOCIE1>': sexe_associe_1sas2,
                        '<PRENOMASSOCIE1>': prenom_associe_1sas2,
                        '<NOMASSOCIE1>': nom_associe_1sas2.upper(),
                        '<DATEDENAISSANCEASSOCIE1<': date_associe_1sas2,
                        '<LIEUDENAISSANCEASSOCIE1>': lieu_naissance_associe_1sas2,
                        '<CODEPOSTALASSOCIE1>': code_postal_associe_1sas2,
                        '<NATIONALITEASSOCIE1>': nationalite_associe_1sas2,
                        '<ADRESSEDOMICILEASSOCIE1>': adresse_associe_1sas2,
                        '<APPORTASSOCIE1>' : apport_associe_1sas2,
                        '<SEXEASSOCIE2>': sexe_associe_2sas2,
                        '<PRENOMASSOCIE2>': prenom_associe_2sas2,
                        '<NOMASSOCIE2>': nom_associe_2sas2.upper(),
                        '<DATEDENAISSANCEASSOCIE2<': date_associe_2sas2,
                        '<LIEUDENAISSANCEASSOCIE2>': lieu_naissance_associe_2sas2,
                        '<CODEPOSTALASSOCIE2>': code_postal_associe_2sas2,
                        '<NATIONALITEASSOCIE2>': nationalite_associe_2sas2,
                        '<ADRESSEDOMICILEASSOCIE2>': adresse_associe_2sas2,
                        '<APPORTASSOCIE2>' : apport_associe_2sas2,
                        '<ACTIVITE>': activitesas2,
                        '<CAPITALENCHIFFRE>': capitalsas2,
                        '<CAPITALENLETTRE>' : num2words(capitalsas2, lang='fr'),
                        '<ANNEEPREMIERBILAN>': annee_bilansas2,
                        '<VILLESIEGESOCIAL>' : ville_siege_socialsas2,
                        '<PARAPHESASSOCIE1>' : get_initials(prenom_associe_1sas2,nom_associe_1sas2),
                        '<PARAPHESASSOCIE2>' : get_initials(prenom_associe_2sas2,nom_associe_2sas2),
                        '<DATE>' : date_aujourd_hui,
                        '<PARENT1>' : pere_associe_1sas2.upper(),
                        '<PARENT2>' : mere_associe_1sas2.upper(),

                    }
                    
                    status_file = remplacer_valeurs('societe/sas2/Statuts SAS 2.docx', remplacements_sas2,"Status "+denominationsas2)
                    liste_action_file = remplacer_valeurs('societe/sas2/Liste souscripteurs actions SAS 2.docx', remplacements_sas2,"Liste souscripteurs "+denominationsas2)
                    dnc_file = remplacer_valeurs('societe/DNC.docx', remplacements_sas2,"DNC "+denominationsas2)

                    infosas2 = f"""
Nom : {nomsas2}
Email : {emailsas2}
Tel : {phonesas2}
"""

                    if status_file:
                        pieces_jointes = [status_file, liste_action_file, dnc_file]
                        email_sender.envoyer_email(destinataire, sujet+ denominationsas2, message+infosas2, pieces_jointes)
                        st.success("Votre formulaire a bien été saisi et envoyé par e-mail.")
                    else:
                        st.error("Problème dans la mise à jour du fichier.")
                else:
                    st.error("Veuillez remplir tous les champs.")

    elif selected_option == "SAS (3 actionnaires)":
        st.header("Société par actions simplifiée (3 actionnaires)")
        with st.form(key='sas3', clear_on_submit=True):
            col1, col2, col3 = st.columns(3)
            with col1:
                denominationsas3 = st.text_input("Nom de la société : ", key="denominationsas3")
                capitalsas3 = st.text_input("Capital social : ", key="capitalsas3")

                st.markdown("**Président associé 1**")
                sexe_associe_1sas3 = st.radio("Sexe ", key="sexe_associe_1sas3", options=["Monsieur", "Madame"])
                prenom_associe_1sas3 = st.text_input("Prénom  : ", key="prenom_associe_1sas3")
                nom_associe_1sas3 = st.text_input("Nom  : ", key="nom_associe_1sas3")
                date_associe_1sas3 = st.text_input("Date de naissance  (JJ/MM/AAAA) : ", key="date_associe_1sas3")
                lieu_naissance_associe_1sas3 = st.text_input("Lieu de naissance :", key="lieu_naissance_associe_1sas3")
                code_postal_associe_1sas3 = st.text_input("Departement de naissance :", key="code_postal_associe_1sas3")
                nationalite_associe_1sas3 = st.text_input("Nationalité :", key="nationalite_associe_1sas3")
                adresse_associe_1sas3 = st.text_input("Adresse de domicile (adresse complète) :", key="adresse_associe_1sas3")
                apport_associe_1sas3 = st.text_input("Apport :", key="apport_associe_1sas3")
                pere_associe_1sas3 = st.text_input("Nom et prénom du père :", key="pere_associe_1sas3")
                mere_associe_1sas3 = st.text_input("Nom et prénom de la mère :", key="mere_associe_1sas3")

                st.markdown('#')
                st.markdown('Vos informations')


                nomsas3 = st.text_input("Nom et prénom :", key="nomsas3")
                emailsas3 = st.text_input("Email :", key="emailsas3")
                phonesas3 = st.text_input("Telephone :", key="phonesas3")


            with col2:
                activitesas3 = st.text_input("Activité de la société : ", key="activitesas3")
                siege_socialsas3 = st.text_input("Siège social (adresse complète) : ", key="siege_socialsas3")

                st.markdown("**Associé 2**")
                sexe_associe_2sas3 = st.radio("Sexe ", key="sexe_associe_2sas3", options=["Monsieur", "Madame"])
                prenom_associe_2sas3 = st.text_input("Prénom :", key="prenom_associe_2sas3")
                nom_associe_2sas3 = st.text_input("Nom :", key="nom_associe_2sas3")
                date_associe_2sas3 = st.text_input("Date de naissance (JJ/MM/AAAA) : ", key="date_associe_2sas3")
                lieu_naissance_associe_2sas3 = st.text_input("Lieu de naissance :", key="lieu_naissance_associe_2sas3")
                code_postal_associe_2sas3 = st.text_input("Departement de naissance :", key="code_postal_associe_2sas3")
                nationalite_associe_2sas3 = st.text_input("Nationalité :", key="nationalite_associe_2sas3")
                adresse_associe_2sas3 = st.text_input("Adresse de domicile (adresse complète) :", key="adresse_associe_2sas3")
                apport_associe_2sas3 = st.text_input("Apport :", key="apport_associe_2sas3")

            with col3:
                annee_bilansas3 = st.text_input("Année de clôture du premier bilan :", key="annee_bilansas3")
                ville_siege_socialsas3 = st.text_input("Ville du siège social : ", key="ville_siege_socialsas3")

                st.markdown("**Associé 3**")
                sexe_associe_3sas3 = st.radio("Sexe ", key="sexe_associe_3sas3", options=["Monsieur", "Madame"])
                prenom_associe_3sas3 = st.text_input("Prénom :", key="prenom_associe_3sas3")
                nom_associe_3sas3 = st.text_input("Nom :", key="nom_associe_3sas3")
                date_associe_3sas3 = st.text_input("Date de naissance (JJ/MM/AAAA) : ", key="date_associe_3sas3")
                lieu_naissance_associe_3sas3 = st.text_input("Lieu de naissance :", key="lieu_naissance_associe_3sas3")
                code_postal_associe_3sas3 = st.text_input("Departement de naissance :", key="code_postal_associe_3sas3")
                nationalite_associe_3sas3 = st.text_input("Nationalité :", key="nationalite_associe_3sas3")
                adresse_associe_3sas3 = st.text_input("Adresse de domicile (adresse complète) :", key="adresse_associe_3sas3")
                apport_associe_3sas3 = st.text_input("Apport :", key="apport_associe_3sas3")


            submittedsas3 = st.form_submit_button("Valider")

            if submittedsas3:
                if all([denominationsas3, capitalsas3, activitesas3, siege_socialsas3, prenom_associe_1sas3, nom_associe_1sas3, date_associe_1sas3,
                        lieu_naissance_associe_1sas3, code_postal_associe_1sas3, nationalite_associe_1sas3, adresse_associe_1sas3,
                        annee_bilansas3, prenom_associe_2sas3, nom_associe_2sas3, date_associe_2sas3, lieu_naissance_associe_2sas3,
                        code_postal_associe_2sas3, nationalite_associe_2sas3, adresse_associe_2sas3, apport_associe_1sas3,apport_associe_2sas3,
                        ville_siege_socialsas3,pere_associe_1sas3, mere_associe_1sas3,prenom_associe_3sas3,nom_associe_3sas3,date_associe_3sas3,
                        lieu_naissance_associe_3sas3,code_postal_associe_3sas3,nationalite_associe_3sas3,adresse_associe_3sas3,apport_associe_3sas3]):
                    
                    remplacements_sas3 = {
                        '<CAPITAL>': capitalsas3,
                        '<SIEGESOCIAL>': siege_socialsas3,
                        '<DENOMINATION>': denominationsas3.upper(),
                        '<SEXEASSOCIE1>': sexe_associe_1sas3,
                        '<PRENOMASSOCIE1>': prenom_associe_1sas3,
                        '<NOMASSOCIE1>': nom_associe_1sas3.upper(),
                        '<DATEDENAISSANCEASSOCIE1<': date_associe_1sas3,
                        '<LIEUDENAISSANCEASSOCIE1>': lieu_naissance_associe_1sas3,
                        '<CODEPOSTALASSOCIE1>': code_postal_associe_1sas3,
                        '<NATIONALITEASSOCIE1>': nationalite_associe_1sas3,
                        '<ADRESSEDOMICILEASSOCIE1>': adresse_associe_1sas3,
                        '<APPORTASSOCIE1>' : apport_associe_1sas3,
                        '<SEXEASSOCIE2>': sexe_associe_2sas3,
                        '<PRENOMASSOCIE2>': prenom_associe_2sas3,
                        '<NOMASSOCIE2>': nom_associe_2sas3.upper(),
                        '<DATEDENAISSANCEASSOCIE2<': date_associe_2sas3,
                        '<LIEUDENAISSANCEASSOCIE2>': lieu_naissance_associe_2sas3,
                        '<CODEPOSTALASSOCIE2>': code_postal_associe_2sas3,
                        '<NATIONALITEASSOCIE2>': nationalite_associe_2sas3,
                        '<ADRESSEDOMICILEASSOCIE2>': adresse_associe_2sas3,
                        '<APPORTASSOCIE2>' : apport_associe_2sas3,
                        '<ACTIVITE>': activitesas3,
                        '<CAPITALENCHIFFRE>': capitalsas3,
                        '<CAPITALENLETTRE>' : num2words(capitalsas3, lang='fr'),
                        '<ANNEEPREMIERBILAN>': annee_bilansas3,
                        '<VILLESIEGESOCIAL>' : ville_siege_socialsas3,
                        '<SEXEASSOCIE3>': sexe_associe_3sas3,
                        '<PRENOMASSOCIE3>': prenom_associe_3sas3,
                        '<NOMASSOCIE3>': nom_associe_3sas3.upper(),
                        '<DATEDENAISSANCEASSOCIE3<': date_associe_3sas3,
                        '<LIEUDENAISSANCEASSOCIE3>': lieu_naissance_associe_3sas3,
                        '<CODEPOSTALASSOCIE3>': code_postal_associe_3sas3,
                        '<NATIONALITEASSOCIE3>': nationalite_associe_3sas3,
                        '<ADRESSEDOMICILEASSOCIE3>': adresse_associe_3sas3,
                        '<APPORTASSOCIE3>' : apport_associe_3sas3,
                        '<PARAPHESASSOCIE1>' : get_initials(prenom_associe_1sas3,nom_associe_1sas3),
                        '<PARAPHESASSOCIE2>' : get_initials(prenom_associe_2sas3,nom_associe_2sas3),
                        '<PARAPHESASSOCIE3>' : get_initials(prenom_associe_3sas3,nom_associe_3sas3),
                        '<PARENT1>' : pere_associe_1sas3.upper(),
                        '<PARENT2>' : mere_associe_1sas3.upper(),
                        '<DATE>' : date_aujourd_hui
                    }
                    
                    status_file = remplacer_valeurs('societe/sas3/Statuts SAS 3.docx', remplacements_sas3,"Status "+denominationsas3)
                    liste_action_file = remplacer_valeurs('societe/sas3/Liste souscripteurs actions SAS 3.docx', remplacements_sas3,"Liste souscripteurs "+denominationsas3)
                    dnc_file = remplacer_valeurs('societe/DNC.docx', remplacements_sas3,"DNC "+denominationsas3)

                    infosas3 = f"""
Nom : {nomsas3}
Email : {emailsas3}
Tel : {phonesas3}
"""
                    
                    if status_file:
                        pieces_jointes = [status_file, liste_action_file, dnc_file]
                        email_sender.envoyer_email(destinataire, sujet+ denominationsas3, message+infosas3, pieces_jointes)
                        st.success("Votre formulaire a bien été saisi et envoyé par e-mail.")
                    else:
                        st.error("Problème dans la mise à jour du fichier.")
                else:
                    st.error("Veuillez remplir tous les champs.")

    elif selected_option == "SASU":
        st.header("Société par actions simplifiée unipersonnelle")
        with st.form(key='sasu', clear_on_submit=True):
            col1, col2 = st.columns(2)
            with col1:
                
                sexe_associe_1sasu = st.radio("Sexe", key="sexe_associe_1sasu", options=["Monsieur", "Madame"])
                prenom_associe_1sasu = st.text_input("Prénom du Président :", key="prenom_associe_1sasu")
                nom_associe_1sasu = st.text_input("Nom du Président :", key="nom_associe_1sasu")
                date_associe_1sasu = st.text_input("Date de naissance (JJ/MM/AAAA) :", key="date_associe_1sasu")
                lieu_naissance_associe_1sasu = st.text_input("Lieu de naissance :", key="lieu_naissance_associe_1sasu")
                code_postal_associe_1sasu = st.text_input("Departement de naissance :", key="code_postal_associe_1sasu")
                nationalite_associe_1sasu = st.text_input("Nationalité :", key="nationalite_associe_1sasu")
                apport_associe_1sasu = st.text_input("Apport :", key="apport_associe_1sasu")
                pere_associe_1sasu = st.text_input("Nom et prénom du père :", key="pere_associe_1sasu")
                mere_associe_1sasu = st.text_input("Nom et prénom de la mère :", key="mere_associe_1sasu")
                adresse_associe_1sasu = st.text_input("Adresse de domicile (adresse complète) :", key="adresse_associe_1sasu")

                st.markdown('#')
                st.markdown('Vos informations')


                nomsasu = st.text_input("Nom et prénom :", key="nomsasu")
                emailsasu = st.text_input("Email :", key="emailsasu")
                phonesasu = st.text_input("Telephone :", key="phonesasu")

                
            with col2:
                denominationsasu = st.text_input("Nom de la société :", key="denominationsasu")
                capitalsasu = st.text_input("Capital social :", key="capitalsasu")
                activitesasu = st.text_input("Activité de la société :", key="activitesasu")
                siege_socialsasu = st.text_input("Siège social (adresse complète) : ", key="siege_socialsasu")
                ville_siege_socialsasu = st.text_input("Ville du siège social :", key="ville_siege_socialsasu")
                annee_bilansasu = st.text_input("Année de clôture du premier bilan :", key="annee_bilansasu")

                


            submittedsasu = st.form_submit_button("Valider")

            if submittedsasu:
                if all([denominationsasu, capitalsasu, activitesasu, siege_socialsasu, prenom_associe_1sasu, nom_associe_1sasu, date_associe_1sasu,
                        lieu_naissance_associe_1sasu, code_postal_associe_1sasu, nationalite_associe_1sasu, adresse_associe_1sasu,
                        annee_bilansasu, apport_associe_1sasu,ville_siege_socialsasu,pere_associe_1sasu, mere_associe_1sasu]):
                    
                    remplacements_sasu = {
                        '<CAPITAL>': capitalsasu,
                        '<SIEGESOCIAL>': siege_socialsasu,
                        '<DENOMINATION>': denominationsasu.upper(),
                        '<SEXEASSOCIE1>': sexe_associe_1sasu,
                        '<PRENOMASSOCIE1>': prenom_associe_1sasu,
                        '<NOMASSOCIE1>': nom_associe_1sasu.upper(),
                        '<DATEDENAISSANCEASSOCIE1<': date_associe_1sasu,
                        '<LIEUDENAISSANCEASSOCIE1>': lieu_naissance_associe_1sasu,
                        '<CODEPOSTALASSOCIE1>': code_postal_associe_1sasu,
                        '<NATIONALITEASSOCIE1>': nationalite_associe_1sasu,
                        '<ADRESSEDOMICILEASSOCIE1>': adresse_associe_1sasu,
                        '<APPORTASSOCIE1>' : apport_associe_1sasu,
                        '<ACTIVITE>': activitesasu,
                        '<CAPITALENCHIFFRE>': capitalsasu,
                        '<CAPITALENLETTRE>' : num2words(capitalsasu, lang='fr'),
                        '<ANNEEPREMIERBILAN>': annee_bilansasu,
                        '<VILLESIEGESOCIAL>' : ville_siege_socialsasu,
                        '<PARAPHESASSOCIE1>' : get_initials(prenom_associe_1sasu,nom_associe_1sasu),
                        '<PARENT1>' : pere_associe_1sasu.upper(),
                        '<PARENT2>' : mere_associe_1sasu.upper(),
                        '<DATE>' : date_aujourd_hui
                    }
                    
                    status_file = remplacer_valeurs('societe/sasu/Statuts_SASU.docx', remplacements_sasu,"Status "+denominationsasu)
                    liste_action_file = remplacer_valeurs('societe/sasu/Liste_souscripteurs_actions_SASU.docx', remplacements_sasu,"Liste souscripteurs "+denominationsasu)
                    dnc_file = remplacer_valeurs('societe/DNC.docx', remplacements_sasu,"DNC "+denominationsasu)

                    infosasu = f"""
Nom : {nomsasu}
Email : {emailsasu}
Tel : {phonesasu}
"""
                    
                    if status_file:
                        pieces_jointes = [status_file, liste_action_file, dnc_file]
                        email_sender.envoyer_email(destinataire, sujet+ denominationsasu, message+infosasu, pieces_jointes)
                        st.success("Votre formulaire a bien été saisi et envoyé par e-mail.")
                    else:
                        st.error("Problème dans la mise à jour du fichier.")
                else:
                    st.error("Veuillez remplir tous les champs.")

    st.markdown("<h2 style='text-align: center;'>Contact</h2>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center;'>Adresse : 75 rue de la Tombe Issoire, 75014 Paris</p>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center;'>Tel : 06.52.24.73.49</p>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center;'>Email : juridique.a2f@gmail.com</p>", unsafe_allow_html=True)

def get_initials(prenom, nom):
    return f"{prenom[0].upper()}{nom[0].upper()}"

def remplacer_valeurs(docx_file, remplacements, nomfichier):
    doc = Document(docx_file)

    def remplacer_dans_paragraph(paragraph, remplacements):
        for old_value, new_value in remplacements.items():
            if old_value in paragraph.text:
                for run in paragraph.runs:
                    if old_value in run.text:
                        run.text = run.text.replace(old_value, new_value)

    def remplacer_dans_table(table, remplacements):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    remplacer_dans_paragraph(paragraph, remplacements)

    # Remplacement dans les paragraphes principaux
    for paragraph in doc.paragraphs:
        remplacer_dans_paragraph(paragraph, remplacements)

    # Remplacement dans les tableaux du corps principal du document
    for table in doc.tables:
        remplacer_dans_table(table, remplacements)

    # Remplacement dans les sections (en-têtes et pieds de page)
    for section in doc.sections:
        # Remplacement dans les en-têtes
        for header in section.header.paragraphs:
            remplacer_dans_paragraph(header, remplacements)
        for table in section.header.tables:
            remplacer_dans_table(table, remplacements)

        # Remplacement dans les pieds de page
        for footer in section.footer.paragraphs:
            remplacer_dans_paragraph(footer, remplacements)
        for table in section.footer.tables:
            remplacer_dans_table(table, remplacements)

    # Sauvegarde du nouveau document
    new_docx_file = nomfichier + ".docx"
    doc.save(new_docx_file)

    return new_docx_file





if __name__ == "__main__":
    main()
